import praw
import time
import os
from docx import *
from common import *
from docx import Document
from playsound import playsound

reddit = praw.Reddit('bot1')

subreddit = reddit.subreddit("all")

document = Document()
document.save("posts.docx")

#counter of how many times program has run
x = 0

#counter of how many times word document was written to
y = 0

#if you do not want the bot to use posts from a particular subreddit, add the subreddit name in the quotes
#Example line: subs_to_ignore = ['funny' , 'pics']
subs_to_ignore = []

while True:
    for submission in subreddit.new():
        sub = submission.subreddit

        if sub in subs_to_ignore:
            print("User has opted to skip this subreddit... Continuing.")

        else:
            if sub.over18:
                continue
                print("NSFW")

            elif sub == "politics":
                continue
                print("politics")


            elif submission.is_self == True:
                writeData = f"Post by: {submission.author}\n\n"
                p = document.add_paragraph(writeData)
                p.add_run(f"Title: {submission.title}\n\n").bold = True
                p.add_run(f"{submission.selftext}\nr/{sub}")
                document.add_paragraph("-------------------------------")
                document.save("posts.docx")
                #print("This is a self post")


                y+=1


            elif "reddit.com/r" not in submission.url:
                #print(submission.url)
                postUrl = submission.url
                formattedLink = f"www.reddit.com{submission.permalink}"
                writeData = f"This is an image post. \n\n url here: {formattedLink[:-1]}"
                document.add_paragraph(writeData)
                document.add_paragraph("-------------------------------")
                document.save("posts.docx")

                y+=1

            elif "snap" or "snapchat" in submission.title:
                continue



        x+=1
        if x % 100 == 0:
            print(x)
            time.sleep(3)
            clear()

        elif x == 500:
            playsound('alert.wav')
            print(f"Wrote to word document {y} times")
            exit(100)
